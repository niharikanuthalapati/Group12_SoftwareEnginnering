import os
import tempfile

from io import BytesIO
import matplotlib
matplotlib.use('Agg')  # Set the backend before importing pyplot
import matplotlib.pyplot as plt
from matplotlib.figure import Figure
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls
from docx.shared import Inches, Pt
from reportlab.lib import colors
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle
from reportlab.pdfgen import canvas
from reportlab.platypus import Paragraph, Table, TableStyle
from django.core.files.base import ContentFile
from django.core.files.storage import default_storage


class ReportGenerator:

    def __init__(self, file_output, ui_feedbacks, review_feedbacks, colorOptions):
        self.file_output = file_output
        self.ui_feedbacks = ui_feedbacks
        self.review_feedbacks = review_feedbacks
        self.colorOptions = colorOptions

    def create_pie_chart(self):
        sentiment_summary = self.file_output.sentiment_summary
        labels = sentiment_summary['labels']
        data = sentiment_summary['datasets'][0]['data']
        colors = self.colorOptions
        
        fig, ax = plt.subplots()
        ax.pie(data, labels=labels, colors=colors, autopct='%1.1f%%')
        ax.set_title("Pie Chart - Sentiment Distribution")
        plt.axis('equal')  # Equal aspect ratio ensures that pie is drawn as a circle.

        # Save pie chart to a bytes buffer
        buf = BytesIO()
        plt.savefig(buf, format='png')
        plt.close(fig)
        return buf.getvalue()

    def create_bar_chart(self):
        sentiment_summary = self.file_output.sentiment_summary
        labels = sentiment_summary['labels']
        data = sentiment_summary['datasets'][0]['data']
        
        fig, ax = plt.subplots()
        ax.bar(labels, data, color=self.colorOptions)
        ax.set_title("Bar Chart - Sentiment Counts")
        
        # Save bar chart to a bytes buffer
        buf = BytesIO()
        plt.savefig(buf, format='png')
        plt.close(fig)
        return buf.getvalue()
    
    def save_chart_to_temp_file(self, chart_data, chart_type='png'):
        # Create a temporary file
        file_descriptor, file_path = tempfile.mkstemp(suffix=f'.{chart_type}')
        # Write chart data to it
        with os.fdopen(file_descriptor, 'wb') as temp_file:
            temp_file.write(chart_data)
        return file_path

    def generate_pdf_report(self):
        # Create charts
        pie_chart_data = self.create_pie_chart()
        bar_chart_data = self.create_bar_chart()

        # Save charts to temporary files
        pie_chart_path = self.save_chart_to_temp_file(pie_chart_data)
        bar_chart_path = self.save_chart_to_temp_file(bar_chart_data)

        # PDF generation setup
        pdf_filename = f'report_{self.file_output.review_file.id}.pdf'
        buffer = BytesIO()
        c = canvas.Canvas(buffer, pagesize=A4)
        width, height = A4  # Width and height of the A4 page

        # Starting Y position (from top of the page, moving down)
        y_position = height - 30

        # Draw Report Title
        c.setFont("Helvetica-Bold", 16)
        c.drawCentredString(width / 2, y_position, "Report")
        y_position -= 30  # Move down for the next element

        style = ParagraphStyle(name='Normal', fontName='Helvetica', fontSize=12, leading=15)

        # Create a paragraph with the review text
        review_text_paragraph = Paragraph(self.file_output.review_text, style)

        # Calculate the height the paragraph will take and update y_position accordingly
        text_width, text_height = review_text_paragraph.wrapOn(c, width - 144, height)  # 72 points margin on each side
        y_position -= text_height + 40  # Subtract the height of the text block plus some padding

        # Draw the paragraph on the canvas
        if y_position < 0:
            y_position = height - text_height - 20

        review_text_paragraph.drawOn(c, 72, y_position)
        # Draw pie chart
        c.drawImage(pie_chart_path, 72, y_position - 200, width=400, height=200)
        y_position -= 220  # Move down for the next element, including some padding
        # Draw bar chart
        c.drawImage(bar_chart_path, 72, y_position - 200, width=400, height=200)
        y_position -= 220  # Move down for the next element, including some padding

        # Add summary table
        y_position = self.add_summary_table_to_pdf(c, (width, y_position))
        # Assuming add_summary_table_to_pdf adjusts y_position internally

        # Add a new page for the feedback tables
        c.showPage()
        page_width, page_height = A4  # Assuming A4 page size

        # Starting Y position for the new page (from the top, going down)
        y_position = page_height - 30  # Start below the top margin

        # Draw the UI Feedback table
        y_position = self.add_feedback_table_to_pdf(c, self.ui_feedbacks, 'User Interface Feedback', y_position, page_width, False)

        # Add some space before the next table
        y_position -= 20

        # Draw the Review Feedback table
        y_position = self.add_feedback_table_to_pdf(c, self.review_feedbacks, 'Review Feedback', y_position, page_width,True)

        # Finalize the PDF and save
        c.showPage()
        c.save()

        # Get the buffer's content as a Django friendly ContentFile
        pdf_content = ContentFile(buffer.getvalue())
        # Save the ContentFile to the media directory using default_storage
        pdf_path = default_storage.save(pdf_filename, pdf_content)

        # Make sure to close the buffer and clean up temporary
        buffer.close()
        # Clean up temporary files
        os.remove(pie_chart_path)
        os.remove(bar_chart_path)
        return pdf_path

    def add_summary_table_to_pdf(self, c, page_dimensions):

        width, y_position = page_dimensions
        sentiment_summary = self.file_output.sentiment_summary
        data = [['Review Type', 'Total', 'Percentage']]

        c.setFont("Helvetica-Bold", 14)
        c.drawString(72, y_position, 'Summary Table')
        y_position -= 20  # Space after title
        
        for label, count, percent in zip(sentiment_summary['labels'], sentiment_summary['datasets'][0]['data'], sentiment_summary['datasets'][0]['percentages']):
            data.append([label, str(count), f"{percent:.2f}%"])
        
        table = Table(data, colWidths=[width * 0.27] * 3)
        table.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
            ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
            ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
            ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
            ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
            ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
            ('GRID', (0, 0), (-1, -1), 1, colors.black),
        ]))
        
        table_width, table_height = table.wrapOn(c, width, y_position)
        table.drawOn(c, 72, y_position - table_height)

        return y_position - table_height - 10  # Return updated y_position

    def add_feedback_table_to_pdf(self, c, feedbacks, title, y_position, page_width, include_rating=False):
        # Add a title for the feedback table
        c.setFont("Helvetica-Bold", 14)
        c.drawString(72, y_position, title)
        y_position -= 20  # Space after title

        # Determine the column widths and headers based on whether rating is included
        if include_rating:
            colWidths = [page_width * 0.266667] * 3
            data = [['Comment', 'Created Date', 'Rating']]
        else:
            colWidths = [page_width * 0.4] * 2
            data = [['Comment', 'Created Date']]

        # Populate table data
        for feedback in feedbacks:
            row = [feedback.comment, feedback.created_date.strftime("%Y-%m-%d %H:%M:%S")]
            if include_rating:
                row.append(str(feedback.star_rating))
            data.append(row)

        # Create the table and style it
        table = Table(data, colWidths=colWidths)
        table.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
            ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
            ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
            ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
            ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
            ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
            ('GRID', (0, 0), (-1, -1), 1, colors.black),
        ]))

        # Calculate the height the table will take and draw the table
        table.wrapOn(c, page_width, y_position)
        table_height = table.wrap(page_width, y_position)[1]
        table.drawOn(c, 72, y_position - table_height - 20)  # Extra 20 for spacing

        # Return the new y_position, above the table and with some padding
        return y_position - table_height - 30

    def set_cell_style(self, cell, bold=False, background_color=None):
        paragraphs = cell.paragraphs
        for paragraph in paragraphs:
            for run in paragraph.runs:
                run.font.size = Pt(11)
                if bold:
                    run.font.bold = True
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        if background_color:
            shading_elm = parse_xml(r'<w:shd {} w:fill="{}"/>'.format(nsdecls('w'), background_color))
            cell._tc.get_or_add_tcPr().append(shading_elm)
    
    def set_cell_background_color(self, cell, rgb_color):
        shading_elm = OxmlElement('w:shd')
        shading_elm.set(qn('w:fill'), rgb_color)
        cell._tc.get_or_add_tcPr().append(shading_elm)

    def generate_docx_report(self):
        # Initialize a Document object
        doc = Document()
        doc.add_heading('Report', level=0)

        # Add review text as a paragraph
        paragraph = doc.add_paragraph(self.file_output.review_text)
        paragraph.style.font.size = Pt(12)

        # Add charts as images
        pie_chart_image = self.create_pie_chart()
        bar_chart_image = self.create_bar_chart()

        # Save charts to temporary files and add to doc
        pie_chart_path = self.save_chart_to_temp_file(pie_chart_image, 'png')
        bar_chart_path = self.save_chart_to_temp_file(bar_chart_image, 'png')

        doc.add_picture(pie_chart_path, width=Inches(6))
        doc.add_picture(bar_chart_path, width=Inches(6))

        # Add sentiment summary table
        self.add_summary_table_to_docx(doc, self.file_output.sentiment_summary)

        # Add feedback tables on a new page
        doc.add_page_break()
        self.add_feedback_table_to_docx(doc, self.ui_feedbacks, 'User Interface Feedback', include_rating=False)
        self.add_feedback_table_to_docx(doc, self.review_feedbacks, 'Review Feedback', include_rating=True)

        # Save the docx file to a byte stream
        docx_stream = BytesIO()
        doc.save(docx_stream)
        docx_stream.seek(0)

        # Prepare the ContentFile to be saved
        docx_content = ContentFile(docx_stream.read(), name=f'report_{self.file_output.review_file.id}.docx')
        
        # Use default_storage to save the ContentFile
        docx_path = default_storage.save(docx_content.name, docx_content)

        # Clean up temporary chart files
        os.remove(pie_chart_path)
        os.remove(bar_chart_path)

        # Return the path where the file was saved
        return docx_path

    def add_summary_table_to_docx(self, doc, sentiment_summary):
        table = doc.add_table(rows=1, cols=3)
        table.style = 'TableGrid'
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Review Type'
        hdr_cells[1].text = 'Total'
        hdr_cells[2].text = 'Percentage'

        # Set header row style
        for cell in hdr_cells:
            self.set_cell_style(cell, bold=True, background_color='D3D3D3')

        # Populate table data
        for label, count, percent in zip(sentiment_summary['labels'], sentiment_summary['datasets'][0]['data'], sentiment_summary['datasets'][0]['percentages']):
            row_cells = table.add_row().cells
            row_cells[0].text = label
            row_cells[1].text = str(count)
            row_cells[2].text = f"{percent:.2f}%"

            # Set data row style
            for cell in row_cells:
                self.set_cell_style(cell)

        return doc

    def add_feedback_table_to_docx(self, doc, feedbacks, title, include_rating=False):
        # Add a title for the feedback table
        doc.add_heading(title, level=2)

        # Determine the number of columns based on whether rating should be included
        num_cols = 3 if include_rating else 2
        # Create a table for feedbacks
        table = doc.add_table(rows=1, cols=num_cols)
        table.style = 'TableGrid'
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Comment'
        hdr_cells[1].text = 'Created Date'
        if include_rating:
            hdr_cells[2].text = 'Rating'

        # Set header row style
        for cell in hdr_cells:
            self.set_cell_style(cell, bold=True, background_color='D3D3D3')

        # Populate table data
        for feedback in feedbacks:
            row_cells = table.add_row().cells
            row_cells[0].text = feedback.comment
            row_cells[1].text = feedback.created_date.strftime("%Y-%m-%d %H:%M:%S")
            if include_rating:
                row_cells[2].text = str(feedback.star_rating) if hasattr(feedback, 'star_rating') else ""

            # Set data row style
            for cell in row_cells:
                self.set_cell_style(cell)

        # Return the document after adding the feedback table
        return doc