import pandas as pd
from textblob import TextBlob
from docx import Document
from matplotlib import pyplot as plt

class DataAnalyzer:
    def __init__(self, file_path):
        self.file_path = file_path
        self.data = pd.read_csv(file_path)
        self.info = {}
        self.pos_percent = 0
        self.neg_percent = 0
        self.neu_percent = 0

    def sentiment_analysis(self):
        def get_sentiment(text):
            blob = TextBlob(text)
            if blob.sentiment.polarity > 0:
                return 'positive'
            elif blob.sentiment.polarity < 0:
                return 'negative'
            else:
                return 'neutral'

        self.data['review_sentiment'] = self.data['review_body'].apply(get_sentiment)
        self.data['headline_sentiment'] = self.data['review_headline'].apply(get_sentiment)

    def analyze_data(self):
        self.sentiment_analysis()

        total_reviews = len(self.data)
        pos_reviews = len(self.data[self.data['review_sentiment'] == 'positive'])
        neg_reviews = len(self.data[self.data['review_sentiment'] == 'negative'])
        neu_reviews = len(self.data[self.data['review_sentiment'] == 'neutral'])

        self.pos_percent = (pos_reviews / total_reviews) * 100
        self.neg_percent = (neg_reviews / total_reviews) * 100
        self.neu_percent = (neu_reviews / total_reviews) * 100

        self.info = {
            'total_reviews': total_reviews,
            'total_positive_reviews': pos_reviews,
            'total_negative_reviews': neg_reviews,
            'total_neutral_reviews': neu_reviews,
        }

    def create_report(self):
        self.analyze_data()

        doc = Document()
        doc.add_heading('Analysis Report', 0)

        # Adding info variable data
        doc.add_heading('Info:', level=1)
        for k, v in self.info.items():
            doc.add_paragraph(f'{k}: {v}')

        # Adding table
        doc.add_heading('Table:', level=1)
        table = doc.add_table(rows=4, cols=3)
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Review Type'
        hdr_cells[1].text = 'Total Number'
        hdr_cells[2].text = 'Percentage'

        row_cells = table.rows[1].cells
        row_cells[0].text = 'Positive'
        row_cells[1].text = str(self.info['total_positive_reviews'])
        row_cells[2].text = f'{self.pos_percent}%'

        row_cells = table.rows[2].cells
        row_cells[0].text = 'Negative'
        row_cells[1].text = str(self.info['total_negative_reviews'])
        row_cells[2].text = f'{self.neg_percent}%'

        row_cells = table.rows[3].cells
        row_cells[0].text = 'Neutral'
        row_cells[1].text = str(self.info['total_neutral_reviews'])
        row_cells[2].text = f'{self.neu_percent}%'

        # Saving the .docx report
        doc.save('report.docx')

        # Creating pie chart
        labels = 'Positive', 'Negative', 'Neutral'
        sizes = [self.pos_percent, self.neg_percent, self.neu_percent]
        colors = ['gold', 'lightcoral', 'lightskyblue']
        explode = (0.1, 0, 0)  # explode 1st slice

        plt.pie(sizes, explode=explode, labels=labels, colors=colors, autopct='%1.1f%%', shadow=True, startangle=140)
        plt.axis('equal')
        plt.savefig('pie_chart.png')

        # Creating bar graph
        objects = ('Positive', 'Negative', 'Neutral')
        y_pos = [0, 1, 2]
        performance = [self.pos_percent, self.neg_percent, self.neu_percent]

        plt.bar(y_pos, performance, align='center', alpha=0.5)
        plt.xticks(y_pos, objects)
        plt.ylabel('Percentage')
        plt.title('Sentiment Analysis')
        plt.savefig('bar_graph.png')

        # Adding graphs to the .pdf report
        from fpdf import FPDF
        pdf = FPDF()
        pdf.add_page()
        pdf.set_font("Arial", size=12)
        pdf.cell(200, 10, txt="Analysis Report", ln=True, align='C')

        pdf.ln(10)  # Add line break

        # Adding info variable data
        pdf.set_font("Arial", size=10)
        pdf.cell(200, 10, txt="Info:", ln=True, align='L')
        for k, v in self.info.items():
            pdf.cell(200, 10, txt=f'{k}: {v}', ln=True, align='L')

        pdf.ln(10)  # Add line break

        # Adding table
        pdf.set_font("Arial", size=10)
        pdf.cell(200, 10, txt="Table:", ln=True, align='L')

        col_widths = [pdf.get_string_width('Review Type') + 6, pdf.get_string_width('Total Number') + 6, pdf.get_string_width('Percentage') + 6]
        pdf.set_fill_color(200, 220, 255)
        pdf.set_text_color(0, 0, 0)
        pdf.set_font("Arial", 'B', 10)

        pdf.cell(col_widths[0], 10, 'Review Type', 1, 0, 'C', 1)
        pdf.cell(col_widths[1], 10, 'Total Number', 1, 0, 'C', 1)
        pdf.cell(col_widths[2], 10, 'Percentage', 1, 1, 'C', 1)

        pdf.set_font("Arial", '', 10)
        pdf.cell(col_widths[0], 10, 'Positive', 1)
        pdf.cell(col_widths[1], 10, str(self.info['total_positive_reviews']), 1)
        pdf.cell(col_widths[2], 10, f'{self.pos_percent}%', 1, 1)

        pdf.cell(col_widths[0], 10, 'Negative', 1)
        pdf.cell(col_widths[1], 10, str(self.info['total_negative_reviews']), 1)
        pdf.cell(col_widths[2], 10, f'{self.neg_percent}%', 1, 1)

        pdf.cell(col_widths[0], 10, 'Neutral', 1)
        pdf.cell(col_widths[1], 10, str(self.info['total_neutral_reviews']), 1)
        pdf.cell(col_widths[2], 10, f'{self.neu_percent}%', 1, 1)

        pdf.ln(10)  # Add line break

        # Adding pie chart
        pdf.image('pie_chart.png', x=10, y=None, w=90)
        pdf.set_x(110)

        # Adding bar graph
        pdf.image('bar_graph.png', x=110, y=None, w=90)

        # Saving the .pdf report
        pdf.output("report.pdf")

if __name__ == '__main__':
    file_path = 'path_to_your_file.csv'
    analyzer = DataAnalyzer(file_path)
    analyzer.create_report()

